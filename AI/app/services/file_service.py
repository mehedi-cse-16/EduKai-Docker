# AutomationCvEmail/app/services/file_service.py
import httpx, io, uuid, os, zipfile
from docx import Document
from pypdf import PdfReader
from fastapi import HTTPException
from PIL import Image

MAX_FILE_SIZE = 10 * 1024 * 1024 
PHOTO_OUTPUT_DIR = "app/static/extracted_photo"
os.makedirs(PHOTO_OUTPUT_DIR, exist_ok=True)

# ASYNC Version 
async def download_file(url: str) -> tuple[bytes, str]:
    async with httpx.AsyncClient() as client:
        return _handle_download(await client.get(url, timeout=30.0), url)

# SYNC Version 
def download_file_sync(url: str) -> tuple[bytes, str]:
    with httpx.Client() as client:
        return _handle_download(client.get(url, timeout=30.0), url)

# Shared Logic to avoid code duplication 
def _handle_download(response, url):
    response.raise_for_status()
    file_size = len(response.content)
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large. Max 10MB.")
    
    filename = url.split("/")[-1].split("?")[0]
    if not filename:
        filename = "temp_cv.pdf"
    return response.content, filename

def extract_text_from_bytes(file_content: bytes, filename: str) -> str:
    file_stream = io.BytesIO(file_content)
    text = ""
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith('.pdf'):
            reader = PdfReader(file_stream)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            if len(text.strip()) < 50:
                raise ValueError("PDF appears to be empty or scanned images. OCR required.")
        
        elif filename_lower.endswith('.docx'):
            doc = Document(file_stream)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text += paragraph.text + "\n"
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format.")
        
        return text.strip()
    
    except Exception as e:
        print(f"Error parsing file {filename}: {str(e)}")
        raise HTTPException(status_code=422, detail=f"Could not parse file: {str(e)}")

def extract_candidate_photo(file_content: bytes, filename: str, base_url: str) -> str | None:
    file_stream = io.BytesIO(file_content)
    filename_lower = filename.lower()
    
    candidate_image_bytes = None
    candidate_image_ext = "jpg"
    
    try:
        # STRATEGY 1: PDF EXTRACTION
        if filename_lower.endswith('.pdf'):
            reader = PdfReader(file_stream)
            if len(reader.pages) > 0:
                page = reader.pages[0] # Only look at page 1
                
                largest_size = 0
                
                for image_file in page.images:
                    img_data = image_file.data
                    if len(img_data) < 2048:  # Filter: Ignore tiny images (icons, bullets) < 2KB
                        continue
                    
                    if len(img_data) > largest_size: # Heuristic: Keep the largest image found
                        largest_size = len(img_data)
                        candidate_image_bytes = img_data
                        candidate_image_ext = image_file.name.split(".")[-1]
        
        # STRATEGY 2: DOCX EXTRACTION 
        elif filename_lower.endswith('.docx'):
            with zipfile.ZipFile(file_stream) as z:
                media_files = [f for f in z.namelist() if f.startswith('word/media/') and f.lower().endswith(('.png', '.jpg', '.jpeg'))]
                
                largest_size = 0
                
                for media_path in media_files:
                    img_data = z.read(media_path)
                    
                    if len(img_data) < 2048:
                        continue
                    
                    if len(img_data) > largest_size:
                        largest_size = len(img_data)
                        candidate_image_bytes = img_data
                        candidate_image_ext = media_path.split(".")[-1]
        
        if candidate_image_bytes:
            unique_name = f"photo_{uuid.uuid4().hex[:8]}.{candidate_image_ext}"
            save_path = os.path.join(PHOTO_OUTPUT_DIR, unique_name)
            
            with open(save_path, "wb") as f:
                f.write(candidate_image_bytes)
            
            print(f"Extracted photo saved: {save_path}")
            if not base_url.endswith("/"): base_url += "/"
            return f"{base_url}static/extracted_photo/{unique_name}"
    
    except Exception as e:
        print(f"Image Extraction Failed for {filename}: {e}")
        return None
    
    return None